from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
import json
import logging
import re
import textwrap
from pathlib import Path
from typing import Any, Protocol
from zipfile import ZIP_DEFLATED, ZipFile

from .models import DocExportConfig, STUDY_EXPORT_FORMATS, STUDY_EXPORT_STYLES


_LOGGER = logging.getLogger(__name__)
_MARKDOWN_ESCAPE_RE = re.compile(r"([\\`*_{}\[\]()#+\-.!|])")
_MAX_TEXT_CHARS = 2000
_MAX_MARKDOWN_CHARS = 120_000
_CONTENT_TYPES = {
    "markdown": "text/markdown; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xmind": "application/vnd.xmind.workbook",
}


@dataclass(slots=True)
class ExportDocument:
    content: bytes
    filename: str
    content_type: str
    markdown: str
    format: str
    style: str


class StudyExportStore(Protocol):
    def list_interactions(self, limit: int = 20) -> list[dict[str, Any]]: ...
    def list_topics(self, limit: int = 100, subject: str | None = None) -> list[dict[str, Any]]: ...
    def list_mastery_overview(self, limit: int = 20) -> list[dict[str, Any]]: ...
    def list_wrong_questions(
        self,
        *,
        limit: int = 20,
        topic_id: str | None = None,
        statuses: tuple[str, ...] = ("active", "retrying", "resolved"),
    ) -> list[dict[str, Any]]: ...


class DocExporter:
    def __init__(
        self,
        store: StudyExportStore,
        *,
        config: DocExportConfig | None = None,
        styles_dir: Path | None = None,
    ) -> None:
        self._validate_store(store)
        self._store = store
        self._config = config or DocExportConfig()
        self._styles_dir = styles_dir or Path(__file__).resolve().parent / "data" / "export_styles"

    @staticmethod
    def _validate_store(store: object) -> None:
        missing = [
            name
            for name in ("list_interactions", "list_topics", "list_mastery_overview", "list_wrong_questions")
            if not callable(getattr(store, name, None))
        ]
        if missing:
            raise TypeError(f"DocExporter store is missing required methods: {', '.join(missing)}")

    def export(
        self,
        *,
        fmt: str = "markdown",
        style: str | None = None,
        title: str | None = None,
        preview_only: bool = False,
        time_range: str | None = None,
        recent_limit: int | None = None,
        topic_ids: list[str] | tuple[str, ...] | None = None,
        **legacy_options: Any,
    ) -> ExportDocument:
        if time_range is None:
            time_range = str(legacy_options.get("range") or "") or None
        export_format = normalize_format(fmt)
        effective_format = "markdown" if preview_only else export_format
        export_style = self.normalize_style(style or self._config.default_style)
        markdown = self.build_markdown(
            title=title,
            style=export_style,
            time_range=time_range,
            recent_limit=recent_limit,
            topic_ids=topic_ids,
        )
        content = markdown.encode("utf-8") if effective_format == "markdown" else self._render(effective_format, markdown)
        return ExportDocument(
            content=content,
            filename=f"{slugify(title or 'study-notes')}.{extension_for_format(effective_format)}",
            content_type=_CONTENT_TYPES[effective_format],
            markdown=markdown,
            format=effective_format,
            style=export_style,
        )

    def build_markdown(
        self,
        *,
        title: str | None = None,
        style: str | None = None,
        time_range: str | None = None,
        recent_limit: int | None = None,
        topic_ids: list[str] | tuple[str, ...] | None = None,
        **legacy_options: Any,
    ) -> str:
        if time_range is None:
            time_range = str(legacy_options.get("range") or "") or None
        export_style = self.normalize_style(style or self._config.default_style)
        style_payload = self.load_style(export_style)
        limit = max(1, min(200, int(recent_limit or 30)))
        topics_limit = max(1, min(5000, int(style_payload.get("topics_limit") or 500)))
        tone = str(style_payload.get("tone") or "").strip()
        requested_topic_ids = _normalized_topic_ids(topic_ids)

        interactions = self._store.list_interactions(limit=limit)
        topics = self._resolve_topics(topics_limit=topics_limit, topic_ids=requested_topic_ids)
        mastery = self._store.list_mastery_overview(limit=topics_limit)
        wrong_questions = self._store.list_wrong_questions(limit=limit)

        now = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
        heading = escape_markdown(title or style_payload.get("title") or "Study Notes", limit=120)
        lines = [
            f"# {heading}",
            "",
            f"- Exported at: `{now}`",
            f"- Style: `{export_style}`",
        ]
        if tone:
            lines.append(f"- Tone: `{escape_markdown(tone, limit=40)}`")
        if time_range:
            lines.append(f"- Range: {escape_markdown(time_range, limit=120)}")
        lines.extend(["", "## Overview", ""])
        lines.append(f"- Recent interactions: {len(interactions)}")
        lines.append(f"- Topics included: {len(topics)}")
        lines.append(f"- Mastery snapshots: {len(mastery)}")
        lines.append(f"- Wrong questions: {len(wrong_questions)}")

        lines.extend(["", "## Recent Interactions", ""])
        if interactions:
            for item in interactions:
                kind = escape_markdown(item.get("kind"), limit=80)
                input_text = escape_markdown(item.get("input_text"), limit=_MAX_TEXT_CHARS)
                output_text = escape_markdown(item.get("output_text"), limit=_MAX_TEXT_CHARS)
                lines.extend([f"### {kind or 'interaction'}", "", f"- Input: {input_text or '-'}", f"- Output: {output_text or '-'}", ""])
        else:
            lines.append("_No recent interactions._")

        lines.extend(["", "## Knowledge Map", ""])
        if topics:
            for topic in topics[:topics_limit]:
                name = escape_markdown(topic.get("name") or topic.get("id"), limit=120)
                topic_id = escape_markdown(topic.get("id"), limit=120)
                subject = escape_markdown(topic.get("subject"), limit=80)
                chapter = escape_markdown(topic.get("chapter"), limit=120)
                lines.append(f"- **{name}** (`{topic_id}`) - {subject or 'general'}{f' / {chapter}' if chapter else ''}")
        else:
            lines.append("_No topics found._")

        lines.extend(["", "## Mastery", ""])
        if mastery:
            for item in mastery[:topics_limit]:
                topic_name = escape_markdown(item.get("topic_name") or item.get("topic_id"), limit=120)
                level = escape_markdown(item.get("level"), limit=80)
                mastery_value = _safe_float(item.get("mastery"))
                lines.append(f"- {topic_name}: {mastery_value:.0%} mastery{f' ({level})' if level else ''}")
        else:
            lines.append("_No mastery data yet._")

        lines.extend(["", "## Wrong Questions", ""])
        if wrong_questions:
            for item in wrong_questions:
                question = item.get("question") if isinstance(item.get("question"), dict) else {}
                text = escape_markdown(question.get("question") or item.get("expected_answer") or item.get("id"), limit=_MAX_TEXT_CHARS)
                error_type = escape_markdown(item.get("error_type"), limit=80)
                status = escape_markdown(item.get("status"), limit=80)
                lines.append(f"- {text or '-'} ({error_type or 'unknown'}, {status or 'active'})")
        else:
            lines.append("_No wrong-question records yet._")

        markdown = "\n".join(lines).strip() + "\n"
        if len(markdown) > _MAX_MARKDOWN_CHARS:
            markdown = markdown[:_MAX_MARKDOWN_CHARS].rstrip() + "\n\n...[export truncated]\n"
        return markdown

    def _resolve_topics(self, *, topics_limit: int, topic_ids: list[str]) -> list[dict[str, Any]]:
        if not topic_ids:
            return self._store.list_topics(limit=topics_limit)

        get_topic = getattr(self._store, "get_topic", None)
        if callable(get_topic):
            topics: list[dict[str, Any]] = []
            for topic_id in topic_ids:
                topic = get_topic(topic_id)
                if isinstance(topic, dict):
                    topics.append(topic)
            return topics

        count_topics = getattr(self._store, "count_topics", None)
        if callable(count_topics):
            try:
                lookup_limit = max(topics_limit, int(count_topics()))
            except (TypeError, ValueError):
                lookup_limit = topics_limit
        else:
            lookup_limit = max(topics_limit, 5000)

        requested = set(topic_ids)
        return [
            item
            for item in self._store.list_topics(limit=max(1, lookup_limit))
            if str(item.get("id") or "") in requested
        ]

    def normalize_style(self, style: str | None) -> str:
        candidate = str(style or "").strip().lower()
        return candidate if candidate in STUDY_EXPORT_STYLES else "neko"

    def load_style(self, style: str) -> dict[str, Any]:
        style_name = self.normalize_style(style)
        path = self._styles_dir / f"{style_name}.json"
        if path.is_file():
            try:
                payload = json.loads(path.read_text(encoding="utf-8"))
            except (OSError, ValueError, TypeError):
                payload = {}
            if isinstance(payload, dict):
                return payload
        return {"title": "Study Notes", "topics_limit": 500}

    def _render(self, fmt: str, markdown: str) -> bytes:
        if fmt == "pdf":
            return self._render_pdf(markdown)
        if fmt == "docx":
            return self._render_docx(markdown)
        if fmt == "xmind":
            if not self._config.xmind_enabled:
                raise ValueError("XMind export is disabled by doc_export.xmind_enabled")
            return self._render_xmind(markdown)
        raise ValueError(f"unsupported export format: {fmt}")

    def _render_pdf(self, markdown: str) -> bytes:
        if self._config.pdf_backend != "reportlab":
            raise ValueError(f"unsupported PDF backend: {self._config.pdf_backend}")
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except ImportError as exc:
            raise RuntimeError("PDF export requires reportlab to be installed") from exc

        output = BytesIO()
        font_name = "Helvetica"
        pdf = canvas.Canvas(output, pagesize=A4)
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception as exc:
            _LOGGER.warning("PDF Unicode font registration failed; Chinese text may render incorrectly: %s", exc)
            font_name = "Helvetica"
        width, height = A4
        x = 48
        y = height - 48
        pdf.setFont(font_name, 10)
        for raw_line in markdown.splitlines() or [""]:
            line = _pdf_safe_text(raw_line)
            wrapped = textwrap.wrap(line, width=92, replace_whitespace=False, drop_whitespace=False) or [""]
            for part in wrapped:
                if y < 48:
                    pdf.showPage()
                    pdf.setFont(font_name, 10)
                    y = height - 48
                pdf.drawString(x, y, safe_utf8_truncate(part, 92 * 4))
                y -= 14
        pdf.save()
        return output.getvalue()

    def _render_docx(self, markdown: str) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX export requires python-docx to be installed") from exc

        document = Document()
        for line in markdown.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _render_xmind(self, markdown: str) -> bytes:
        root_topic = markdown.splitlines()[0].lstrip("# ").strip() if markdown.strip() else "Study Notes"
        children = []
        for line in markdown.splitlines():
            if line.startswith("## "):
                children.append({"id": slugify(line[3:]) or "section", "title": line[3:].strip()})
        content = [
            {
                "id": "study-companion-sheet",
                "title": root_topic or "Study Notes",
                "rootTopic": {
                    "id": "root",
                    "title": root_topic or "Study Notes",
                    "children": {"attached": children[:30]},
                },
            }
        ]
        output = BytesIO()
        with ZipFile(output, "w", ZIP_DEFLATED) as archive:
            archive.writestr("content.json", json.dumps(content, ensure_ascii=False))
            archive.writestr("metadata.json", json.dumps({"creator": "study_companion"}, ensure_ascii=False))
            archive.writestr("manifest.json", json.dumps({"file-entries": {"content.json": {}, "metadata.json": {}}}))
        return output.getvalue()


def normalize_format(value: str | None) -> str:
    fmt = str(value or "markdown").strip().lower()
    if fmt == "md":
        fmt = "markdown"
    if fmt not in STUDY_EXPORT_FORMATS:
        raise ValueError(f"unsupported export format: {fmt}")
    return fmt


def _normalized_topic_ids(topic_ids: list[str] | tuple[str, ...] | None) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for item in topic_ids or []:
        topic_id = str(item).strip()
        if not topic_id or topic_id in seen:
            continue
        result.append(topic_id)
        seen.add(topic_id)
    return result


def extension_for_format(fmt: str) -> str:
    return "md" if fmt == "markdown" else fmt


def escape_markdown(value: object, *, limit: int = _MAX_TEXT_CHARS) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(text) > limit:
        text = text[:limit].rstrip() + "...[truncated]"
    return _MARKDOWN_ESCAPE_RE.sub(r"\\\1", text)


def slugify(value: object) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"[^a-z0-9._-]+", "-", text)
    text = re.sub(r"-{2,}", "-", text).strip("-._")
    return text[:80] or "study-notes"


def _safe_float(value: object) -> float:
    try:
        return float(value)
    except (TypeError, ValueError, OverflowError):
        return 0.0


def _pdf_safe_text(value: object) -> str:
    return str(value or "").replace("\t", "    ")


def safe_utf8_truncate(text: str, max_bytes: int) -> str:
    if max_bytes <= 0:
        return ""
    payload = str(text or "").encode("utf-8")
    if len(payload) <= max_bytes:
        return str(text or "")
    return payload[:max_bytes].decode("utf-8", errors="ignore")


__all__ = [
    "DocExporter",
    "ExportDocument",
    "StudyExportStore",
    "escape_markdown",
    "extension_for_format",
    "normalize_format",
    "safe_utf8_truncate",
    "slugify",
]
